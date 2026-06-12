from __future__ import annotations

from pathlib import Path

from video_transcriber.models import TranscriptResult
from video_transcriber.utils import render_transcript_text


class TranscriptExporter:
    def export_txt(self, result: TranscriptResult, destination: Path, include_timestamps: bool) -> Path:
        text = render_transcript_text(result, include_timestamps=include_timestamps)
        destination.write_text(text, encoding="utf-8")
        return destination

    def export_docx(self, result: TranscriptResult, destination: Path, include_timestamps: bool) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Falta la dependencia 'python-docx'. Ejecuta 'pip install -e .'") from exc

        document = Document()
        document.add_heading("Transcripcion", level=1)
        document.add_paragraph(f"Archivo origen: {result.source_file.name}")
        if result.detected_language:
            document.add_paragraph(f"Idioma detectado: {result.detected_language}")
        if result.duration_seconds is not None:
            document.add_paragraph(f"Duracion aproximada: {int(result.duration_seconds)} segundos")

        document.add_paragraph("")
        for line in render_transcript_text(result, include_timestamps=include_timestamps).splitlines():
            document.add_paragraph(line)

        document.save(destination)
        return destination

    def export_pdf(self, result: TranscriptResult, destination: Path, include_timestamps: bool) -> Path:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("Falta la dependencia 'reportlab'. Ejecuta 'pip install -e .'") from exc

        styles = getSampleStyleSheet()
        story = [
            Paragraph("Transcripcion", styles["Title"]),
            Spacer(1, 12),
            Paragraph(f"Archivo origen: {result.source_file.name}", styles["BodyText"]),
        ]
        if result.detected_language:
            story.append(Paragraph(f"Idioma detectado: {result.detected_language}", styles["BodyText"]))
        if result.duration_seconds is not None:
            story.append(Paragraph(f"Duracion aproximada: {int(result.duration_seconds)} segundos", styles["BodyText"]))

        story.append(Spacer(1, 12))
        for line in render_transcript_text(result, include_timestamps=include_timestamps).splitlines():
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, styles["BodyText"]))
            story.append(Spacer(1, 6))

        document = SimpleDocTemplate(str(destination), pagesize=A4)
        document.build(story)
        return destination
